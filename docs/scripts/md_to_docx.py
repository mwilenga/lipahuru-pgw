#!/usr/bin/env python3
"""Convert MERCHANT_INTEGRATION_GUIDE.md to a formatted Word document."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        if level == 1:
            style.font.size = Pt(22)
        elif level == 2:
            style.font.size = Pt(16)
        else:
            style.font.size = Pt(13)


def add_code_block(document: Document, lines: list[str]) -> None:
    text = "\n".join(lines).rstrip("\n")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    shading = paragraph._element.get_or_add_pPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): "F1F5F9",
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            cell = table.rows[row_idx].cells[col_idx]
            value = row[col_idx] if col_idx < len(row) else ""
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    if row_idx == 0:
                        run.bold = True

    document.add_paragraph()


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_checklist(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(f"☐ {item}")


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line.strip()))


def add_rich_paragraph(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def convert_markdown(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    set_document_defaults(document)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("```"):
            fence = stripped
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(document, code_lines)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_rows = [parse_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(document, table_rows)
            continue

        if stripped.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_lines.append(lines[i].strip()[2:].strip())
                i += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(" ".join(quote_lines))
            run.italic = True
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            continue

        if re.match(r"^- \[[ x]\] ", stripped):
            checklist_items = []
            while i < len(lines) and re.match(r"^- \[[ x]\] ", lines[i].strip()):
                checklist_items.append(re.sub(r"^- \[[ x]\] ", "", lines[i].strip()))
                i += 1
            add_checklist(document, checklist_items)
            continue

        if stripped.startswith("- "):
            bullet_items = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                bullet_items.append(lines[i].strip()[2:].strip())
                i += 1
            add_bullet_list(document, bullet_items)
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("```")
                or nxt.startswith("|")
                or nxt.startswith("- ")
                or nxt.startswith("> ")
                or nxt == "---"
            ):
                break
            paragraph_lines.append(nxt)
            i += 1
        add_rich_paragraph(document, " ".join(paragraph_lines))

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md_file = root / "MERCHANT_INTEGRATION_GUIDE.md"
    out_file = root / "MERCHANT_INTEGRATION_GUIDE.docx"
    if len(sys.argv) > 1:
        md_file = Path(sys.argv[1])
    if len(sys.argv) > 2:
        out_file = Path(sys.argv[2])
    convert_markdown(md_file, out_file)
    print(f"Created {out_file}")
